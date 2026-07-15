"""Convert-to-file tool — convert JSON/CSV/Markdown/text content into a target file
(PDF, DOCX, CSV, JSON, TXT).

Mirrors: typescript/packages/core/tools/convert_to_file/convert_to_file.ts

Markdown -> PDF/DOCX rendering deliberately avoids a headless browser: Markdown is
tokenized with `mistune` (AST mode) and the resulting blocks (headings, paragraphs,
bullet lists) are drawn directly with `reportlab` (platypus flowables) / assembled
with `python-docx`. This mirrors the lightweight-deps approach `extract_from_file`
takes for reading files.
"""
from __future__ import annotations

import base64
import csv
import io
import json
import re
from pathlib import Path
from typing import Any

import mistune
from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, SimpleDocTemplate, Spacer
from reportlab.platypus import Paragraph as PdfParagraph

from matimo.errors import ErrorCode, MatimoError

SOURCE_FORMATS = ("json", "csv", "markdown", "text")
TARGET_FORMATS = ("pdf", "docx", "csv", "json", "txt")
DEFAULT_MAX_CONTENT_LENGTH = 10 * 1024 * 1024  # 10MB

# The only source_format -> target_format pairs this tool supports.
VALID_COMBOS: tuple[tuple[str, str], ...] = (
    ("json", "csv"),
    ("csv", "json"),
    ("markdown", "pdf"),
    ("markdown", "docx"),
    ("text", "docx"),
    ("text", "txt"),
)
VALID_COMBO_KEYS = {f"{s}->{t}" for s, t in VALID_COMBOS}

MIME_TYPES = {
    "csv": "text/csv",
    "json": "application/json",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
}


def _resolve_local_path(file_path: str) -> Path:
    """Expand ~ and resolve relative paths, mirroring the `read` / `extract_from_file` core tools."""
    if file_path.startswith("~"):
        file_path = str(Path.home()) + file_path[1:]
    return Path(file_path).resolve()


# ── CSV helpers (stdlib csv module — mirrors extract_from_file's approach of
#    avoiding a new CSV dependency) ────────────────────────────────────────


def _parse_csv_rows(text: str) -> list[list[str]]:
    reader = csv.reader(io.StringIO(text))
    return [row for row in reader if row]


def _serialize_csv(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    buf = io.StringIO()
    writer = csv.writer(buf, lineterminator="\r\n")
    writer.writerows(rows)
    return buf.getvalue()


def _stringify_csv_value(value: Any) -> str:  # noqa: ANN401
    if value is None:
        return ""
    if isinstance(value, bool | int | float):
        return str(value)
    if isinstance(value, str):
        return value
    return json.dumps(value)


# ── JSON <-> CSV converters ────────────────────────────────────────────────


def _json_to_csv(content: str) -> bytes:
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as exc:
        raise MatimoError(
            "Invalid JSON content", ErrorCode.INVALID_PARAMETER, {"reason": str(exc)}
        ) from exc

    records: list[dict[str, Any]]
    if isinstance(parsed, list):
        records = [item if isinstance(item, dict) else {"value": item} for item in parsed]
    elif isinstance(parsed, dict):
        records = [parsed]
    else:
        raise MatimoError(
            "Unsupported JSON shape for CSV conversion",
            ErrorCode.INVALID_PARAMETER,
            {"reason": "JSON content must be an object or an array of objects/values to convert to CSV"},
        )

    if not records:
        return b""

    columns: list[str] = []
    seen: set[str] = set()
    for record in records:
        for key in record:
            if key not in seen:
                seen.add(key)
                columns.append(key)
    if not columns:
        columns = ["value"]

    rows = [columns, *[[_stringify_csv_value(record.get(col)) for col in columns] for record in records]]
    return _serialize_csv(rows).encode("utf-8")


def _csv_to_json(content: str) -> bytes:
    rows = _parse_csv_rows(content)
    if not rows:
        return b"[]"

    header, *data_rows = rows
    records = []
    for row in data_rows:
        record: dict[str, str] = {}
        for idx, col in enumerate(header):
            key = col or f"column_{idx + 1}"
            record[key] = row[idx] if idx < len(row) else ""
        records.append(record)

    return json.dumps(records, indent=2).encode("utf-8")


# ── Markdown tokenization (shared by PDF and DOCX renderers) ──────────────

MdBlock = dict[str, Any]  # {"kind": "heading"|"paragraph"|"bullet", "level"?: int, "text": str}

_markdown_parser = mistune.create_markdown(renderer=None)


def _flatten_inline(node: dict[str, Any]) -> str:
    """Recursively flatten a mistune AST node into plain text (formatting markers stripped)."""
    children = node.get("children")
    if children:
        return "".join(_flatten_inline(child) for child in children)
    raw = node.get("raw")
    if isinstance(raw, str):
        return raw
    text = node.get("text")
    if isinstance(text, str):
        return text
    return ""


def _normalize_markdown(content: str) -> list[MdBlock]:
    """Parse Markdown into a flat list of headings/paragraphs/bullets.

    Intentionally lightweight — no full CommonMark fidelity (tables, nested
    lists, images are not specially handled), but structure (headings/
    paragraphs/lists) is never silently dropped: unrecognized block types
    fall back to plain paragraphs.
    """
    ast_nodes = _markdown_parser(content)
    # mistune's Markdown.__call__ is typed `str | list[dict[str, Any]]` because a
    # non-None renderer returns rendered output (e.g. HTML) as a string; with
    # renderer=None (set above) it always returns the raw AST node list.
    assert isinstance(ast_nodes, list)  # noqa: S101
    blocks: list[MdBlock] = []
    for node in ast_nodes:
        node_kind = node.get("type")
        if node_kind == "heading":
            level = node.get("attrs", {}).get("level", 1)
            blocks.append({"kind": "heading", "level": level, "text": _flatten_inline(node)})
        elif node_kind == "paragraph":
            blocks.append({"kind": "paragraph", "text": _flatten_inline(node)})
        elif node_kind == "list":
            for item in node.get("children", []):
                blocks.append({"kind": "bullet", "text": _flatten_inline(item)})
        elif node_kind in ("blank_line", "thematic_break"):
            continue
        else:
            text = _flatten_inline(node)
            if text.strip():
                blocks.append({"kind": "paragraph", "text": text})
    return blocks


# ── Markdown -> PDF (reportlab platypus; no headless browser) ─────────────

_PDF_STYLES = getSampleStyleSheet()
_PDF_HEADING_STYLE_NAMES = {
    1: "Heading1",
    2: "Heading2",
    3: "Heading3",
    4: "Heading4",
    5: "Heading5",
    6: "Heading6",
}


def _markdown_to_pdf(content: str) -> bytes:
    blocks = _normalize_markdown(content)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER)
    flow: list[Any] = []
    bullet_items: list[ListItem] = []

    def flush_bullets() -> None:
        if bullet_items:
            flow.append(ListFlowable(list(bullet_items), bulletType="bullet"))
            flow.append(Spacer(1, 6))
            bullet_items.clear()

    for block in blocks:
        if block["kind"] == "bullet":
            bullet_items.append(ListItem(PdfParagraph(block["text"], _PDF_STYLES["Normal"])))
            continue
        flush_bullets()
        if block["kind"] == "heading":
            style_name = _PDF_HEADING_STYLE_NAMES.get(block["level"], "Heading6")
            flow.append(PdfParagraph(block["text"], _PDF_STYLES[style_name]))
        else:
            flow.append(PdfParagraph(block["text"], _PDF_STYLES["Normal"]))
        flow.append(Spacer(1, 6))
    flush_bullets()

    if not flow:
        flow.append(PdfParagraph("", _PDF_STYLES["Normal"]))

    doc.build(flow)
    return buf.getvalue()


# ── Markdown / text -> DOCX (python-docx) ──────────────────────────────────


def _markdown_to_docx(content: str) -> bytes:
    blocks = _normalize_markdown(content)
    document = Document()
    if not blocks:
        document.add_paragraph("")
    for block in blocks:
        if block["kind"] == "heading":
            level = min(max(block["level"], 1), 6)
            document.add_heading(block["text"], level=level)
        elif block["kind"] == "bullet":
            document.add_paragraph(block["text"], style="List Bullet")
        else:
            document.add_paragraph(block["text"])

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _text_to_docx(content: str) -> bytes:
    document = Document()
    # Split the same way the TypeScript executor does (String.split(/\r?\n/)) so
    # both SDKs produce the same paragraph count, including a trailing empty
    # paragraph when content ends with a newline.
    for line in re.split(r"\r?\n", content):
        document.add_paragraph(line)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ── Main entry point ────────────────────────────────────────────────────────


async def run(params: dict[str, Any]) -> dict[str, Any]:
    content: str | None = params.get("content")
    source_format: str | None = params.get("source_format")
    target_format: str | None = params.get("target_format")
    output_path: str | None = params.get("output_path")
    max_content_length = int(params.get("max_content_length") or DEFAULT_MAX_CONTENT_LENGTH)

    if not isinstance(content, str) or len(content) == 0:
        raise MatimoError(
            "Missing required parameter",
            ErrorCode.INVALID_PARAMETER,
            {"reason": "content is required and must be a non-empty string"},
        )

    if source_format not in SOURCE_FORMATS:
        raise MatimoError(
            "Unsupported source_format",
            ErrorCode.INVALID_PARAMETER,
            {"source_format": source_format, "supported": list(SOURCE_FORMATS)},
        )

    if target_format not in TARGET_FORMATS:
        raise MatimoError(
            "Unsupported target_format",
            ErrorCode.INVALID_PARAMETER,
            {"target_format": target_format, "supported": list(TARGET_FORMATS)},
        )

    combo_key = f"{source_format}->{target_format}"
    if combo_key not in VALID_COMBO_KEYS:
        raise MatimoError(
            "Unsupported conversion combination",
            ErrorCode.INVALID_PARAMETER,
            {
                "source_format": source_format,
                "target_format": target_format,
                "valid_combinations": [f"{s}->{t}" for s, t in VALID_COMBOS],
            },
        )

    if len(content) > max_content_length:
        raise MatimoError(
            "Content too large",
            ErrorCode.EXECUTION_FAILED,
            {"size": len(content), "max_content_length": max_content_length},
        )

    if combo_key == "json->csv":
        data = _json_to_csv(content)
    elif combo_key == "csv->json":
        data = _csv_to_json(content)
    elif combo_key == "markdown->pdf":
        data = _markdown_to_pdf(content)
    elif combo_key == "markdown->docx":
        data = _markdown_to_docx(content)
    elif combo_key == "text->docx":
        data = _text_to_docx(content)
    elif combo_key == "text->txt":
        data = content.encode("utf-8")
    else:  # pragma: no cover — unreachable: guarded by VALID_COMBO_KEYS above
        raise MatimoError(
            "Unsupported conversion combination",
            ErrorCode.INVALID_PARAMETER,
            {"source_format": source_format, "target_format": target_format},
        )

    mime_type = MIME_TYPES[target_format]

    if output_path:
        resolved = _resolve_local_path(output_path)
        resolved.parent.mkdir(parents=True, exist_ok=True)
        resolved.write_bytes(data)
        return {
            "success": True,
            "output_path": str(resolved),
            "file_base64": None,
            "mime_type": mime_type,
            "size_bytes": len(data),
        }

    return {
        "success": True,
        "output_path": None,
        "file_base64": base64.b64encode(data).decode("ascii"),
        "mime_type": mime_type,
        "size_bytes": len(data),
    }
